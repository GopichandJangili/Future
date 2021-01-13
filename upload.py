#OS
import os
#Doc
import shutil
from os import *
from os.path import isfile, join
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
import xlrd
import urllib
import urllib.request
import pickle
from waitress import serve
from docx import Document
from docx.shared import Inches
from flask import send_from_directory
#Pandas
import pandas as pd,numpy as np,sqlite3
import pandas as pd,numpy as np,xlrd
from tqdm import tqdm
from math import *
from Process import va,va1,exportToWord
#Flask API
from flask import Flask, render_template,request,send_file,redirect,url_for,after_this_request,send_from_directory,abort,flash
from flask_cors import CORS
from werkzeug.utils import secure_filename
import json,ast
from functools import reduce
from funcs import *
from io import StringIO
import time
from datetime import datetime
import time
from docx.shared import RGBColor,Pt,Inches
from docx_utils.flatten import opc_to_flat_opc
#from docx2pdf import convert
import docx
import hupper
from waitress import serve

def getText(filename):
    doc = Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)




#Only csv files are allowed
ALLOWED_DOCS = {'xlsx'}
ALLOWED_IMAGES = {'png','jpg','jpeg'}
File_Name = {'Template'}
app = Flask(__name__)
app.secret_key = 'random string'

#Give always a new file for export (as we are using send_file)
app.config['SEND_FILE_MAX_AGE_DEFAULT']=0

#Size 16MB
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024
#Cache control
CORS(app)
app.config['TEMPLATES_AUTO_RELOAD'] = True
def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_DOCS


@app.after_request
def add_header(response):
    # response.cache_control.no_store = True
    if 'Cache-Control' not in response.headers:
        response.headers['Cache-Control'] = 'no-store'

    return response


def deletecsvfile(project):
    path=os.getcwd()
    path1=os.path.join(path,'Data')
    path=os.path.join(path1,project)
    pathfile=os.path.join(path,'Updated_rules.csv')
    if os.path.exists(pathfile):
        os.remove(pathfile)
@app.route('/<project>/downloadfile', methods = ['GET','POST'])
def process(project):
    # do what you're doing
    path=os.getcwd()
    path=os.path.join(path,'Data')
    path=os.path.join(path,project)
    file_name = 'Document.docx'
    return send_from_directory(path,file_name, as_attachment=True)

@app.route('/<project>/delete', methods = ['GET','POST'])
def delete_files(project):

    try:
     print('***************************Deleted******************************')
     print(project)
     path=os.getcwd()
     print(path)
     path=os.path.join(path,'Data')
     print(path)
     path=os.path.join(path,project)
     print(path)
     shutil.rmtree(path)
     print('***************************Deleted******************************')
     flash('Project has been deleted')

     if os.path.exists(os.path.join(os.getcwd(),'Data',project+'.db')):
      os.remove(os.path.join(os.getcwd(),'Data',project+'.db'))
     if os.path.exists(os.path.join(os.getcwd(),'Data',project+'.json')):
      os.remove(os.path.join(os.getcwd(),'Data',project+'.json'))
     if os.path.exists(os.path.join(os.getcwd(),'Data',project+'_content.csv')):
      os.remove(os.path.join(os.getcwd(),'Data',project+'_content.csv'))
     if os.path.exists(os.path.join(os.getcwd(),'Data',project+'_sort.json')):
      os.remove(os.path.join(os.getcwd(),'Data',project+'_sort.json'))

     return redirect('/')
    except Exception as e:
     print('*************************Not Deleted**********************************')
     print(str(e))

     flash("Project couldn't be deleted")
     return redirect('/')

@app.route('/<project>/uploaddata', methods = ['GET','POST'])
def upload_filess(project):
    try:
        deletecsvfile(project)
        if request.method == 'POST':

            if 'file' not in request.files:
                flash('No file part')
                return redirect(request.url)
            file = request.files['file']
            if file.filename == '':
                flash('No selected file')
                return redirect(request.url)
            if os.path.isdir(os.path.join('Data',project)) :
                pass
            else:
                path=os.getcwd()
                path=os.path.join(path,'Data')
                path=os.path.join(path,project)
                path1=os.path.join(path,'Data_defination')
                path2=os.path.join(path,'Data_template')
                os.makedirs(path)
                os.makedirs(path1)
                os.makedirs(path2)
            path=os.getcwd()
            path=os.path.join(path,'Data')
            path=os.path.join(path,project)
            path1=os.path.join(path,'Data_defination')
            path2=os.path.join(path,'Data_template')
            #if file and allowed_file(file.filename):
            filelist = [ os.remove(os.path.join(path1,f)) for f in os.listdir(path1) ]
            print(filelist)
            filename = secure_filename(file.filename)
            file.save(os.path.join(path1, filename))
            flash(u'File has been successfully uploaded')

            #else:
            # flash(u'Check file extension','error')
    except Exception as e:
        print(str(e))
        print(1)
        flash(u'Check file extension','error')
    return redirect(url_for('upload_files', project=project))

@app.route('/<project>/downloaddata', methods = ['GET'])
def downloaddef(project):
     path=os.path.join('Data',project,'Data_defination')
     filename=os.listdir(path)[0]
     return send_from_directory(path,filename, as_attachment=True)


@app.route('/<project>/downloadtemplate', methods = ['GET'])
def downloadtemp(project):
     path=os.path.join('Data',project,'Data_template')
     filename=os.listdir(path)[0]
     return send_from_directory(path,filename, as_attachment=True)


@app.route('/<project>/uploaddatatemplate', methods = ['GET','POST'])
def upload_filessss(project):
    try:
        deletecsvfile(project)
        if request.method == 'POST':
            print(os.getcwd())
            path=os.getcwd()
            path=os.path.join(path,'Data')
            path=os.path.join(path,project)
            path1=os.path.join(path,'Data_defination')
            path2=os.path.join(path,'Data_template')
            #if 'file' not in request.files:
            #    flash(u'Please check file extension','error')
            #    return redirect(request.url)
            file = request.files['file']
            if file.filename == '':
                flash('No selected file')
                return redirect(request.url)
            if os.path.isdir(os.path.join('Data',project)) :
                pass
            else:
                os.makedirs(path)
                os.makedirs(path1)
                os.makedirs(path2)

            #if file and allowed_file(file.filename):
            filelist = [ os.remove(os.path.join(path2,f)) for f in os.listdir(path2) ]
            print(filelist)
            filename = secure_filename(file.filename)
            file.save(os.path.join(path2, filename))
            flash(u'File has been successfully uploaded')

            #else:
            #flash(u'Check file extension','error')

    except Exception as e:
        flash(u'Check file extension','error')
    return redirect(url_for('upload_files', project=project))

    #return redirect(f'/{project}',file_list = f_list,file_list1 = f_list1)
@app.route('/<project1>/uploaddataurl', methods = ['GET','POST'])
def upload_filessurl(project1):
        deletecsvfile(project1)
        print(os.getcwd())
        print("server has called api")
        link = request.form['text']
        #https://docs.google.com/spreadsheets/d/1Gztm9o8JEPibPWEDwH54qBG5kwj51ILDOk_dxK6uTSY/export?format=xlsx&gid=1330027783
        print(link)
        try:
            path=os.getcwd()
            path=os.path.join(path,'Data')
            path=os.path.join(path,project1)
            path1=os.path.join(path,'Data_defination')
            path2=os.path.join(path,'Data_template')
            if os.path.isdir(os.path.join('Data',project1)) :
                pass
            else:
                os.makedirs(path)
                if os.path.isdir(path1) :
                    pass
                else:
                    os.makedirs(path1)
                if os.path.isdir(path2) :
                    pass
                else:
                    os.makedirs(path2)
            filelist = [ os.remove(os.path.join(path1,f)) for f in os.listdir(path1) ]
            print(filelist)
            #file_to_write = open("output.pickle", "wb")
            #pickle.dump(a, file_to_write)
            print('1')
            urllib.request.urlretrieve(link,os.path.join(path1,'Data_Defination.xlsx'))
            print('2')
            return redirect(url_for('upload_files', project=project1))
            #return render_template('Lat.html',file_list = f_list,file_list1 = f_list1)
        except Exception as e:
            Exceptionerror="Please check the url once"
            print(Exceptionerror)
            flash(u'Please check the url once','error')
            return redirect(url_for('upload_files', project=project1))
            #return render_template('Lat.html',file_list = f_list,file_list1 = f_list1)
        return redirect(url_for('upload_files', project=project1))
        #return render_template('Lat.html',file_list = f_list,file_list1 = f_list1)
@app.route('/<project>/uploaddatatemplateurl', methods = ['GET','POST'])
def upload_files_template_url(project):
        deletecsvfile(project)
        print(os.getcwd())
        print("server has called api")
        link = request.form['text1']
        print(link)
        try:
            path=os.getcwd()
            path=os.path.join(path,'Data')
            path=os.path.join(path,project)
            path1=os.path.join(path,'Data_defination')
            path2=os.path.join(path,'Data_template')
            if os.path.isdir(os.path.join('Data',project)) :
                pass
            else:
                os.makedirs(path)
                if os.path.isdir(path1) :
                    pass
                else:
                    os.makedirs(path1)
                if os.path.isdir(path2) :
                    pass
                else:
                    os.makedirs(path2)
            filelist = [ os.remove(os.path.join(path2,f)) for f in os.listdir(path2) ]
            print(filelist)
            print('1')
            urllib.request.urlretrieve(link,os.path.join(path2,'Data_Template.xlsx'))
            print('2')
            return redirect(url_for('upload_files', project=project))
        except Exception as e:
            Exceptionerror="Please check the url once"
            print(Exceptionerror)
            flash(u'Please check the url once','error')
            return redirect(url_for('upload_files', project=project))
        return redirect(url_for('upload_files', project=project))
@app.route('/<project>', methods = ['GET','POST'])
def upload_files(project):
    if request.method=='GET':
        try:
            print(os.getcwd())
            path=os.getcwd()
            path=os.path.join(path,'Data')
            path=os.path.join(path,project)
            path1=os.path.join(path,'Data_defination')
            path2=os.path.join(path,'Data_template')
            if os.path.isdir(os.path.join('Data',project)) :
                pass
            else:
                os.makedirs(path)
                if os.path.isdir(path1) :
                    pass
                else:
                    os.makedirs(path1)
                if os.path.isdir(path2) :
                    pass
                else:
                    os.makedirs(path2)
            f_list = [f for f in os.listdir(path1) if isfile(join(path1, f))]
            f_list1 = [f for f in os.listdir(path2) if isfile(join(path2, f))]
            if os.path.exists(os.path.join(os.getcwd(),'Data',project+'_sort.json')):
                with open(os.path.join(os.getcwd(),'Data',project+'_sort.json')) as json_file:
                    json_sort = json.load(json_file)
                sheets_needed=list(json_sort.keys()  )

            else:
                sheets_needed=None

            input_file  = db_read(project,'input',sheets_needed=sheets_needed)

            colummnsList = input_file
            colummnsList = ['{'+ x +'}' for x in colummnsList]

            df = db_read(project,'template')
            df=df[["Sno","Business_Rule","String_Name"]]
            df=df.replace(np.nan,'',regex=True)
            items = df.to_dict('records')



            if not os.path.exists(os.path.join(os.getcwd(),'Data',project+'_sort.json')):
                json_decoded={'All':''}
            with open(os.path.join(os.getcwd(),'Data',project+'_sort.json'),'w') as json_file:
                json.dump(json_decoded, json_file)
            if os.path.exists(os.path.join(os.getcwd(),'Data',project+'.json')):
                with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
                    json_decoded = json.load(json_file)
            else:
                json_decoded={'key':'#sp','header':'','description':'','datadefinition':'','datatemplate':'','rowbinder':'','format':'#word'}

            with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
                json.dump(json_decoded, json_file)
            return render_template("Lat.html",items=items,columns=colummnsList,
                              project=project,thing=json_decoded['key'],
                              datadefinition=json_decoded['datadefinition'],
                              datatemplate=json_decoded['datatemplate'],
                              rbind=json_decoded['rowbinder'],
                              fmat=json_decoded['format'],file_list = f_list,file_list1 = f_list1)
        except Exception as e:
            print(os.getcwd())
            path=os.getcwd()
            path=os.path.join(path,'Data')
            path=os.path.join(path,project)
            path1=os.path.join(path,'Data_defination')
            path2=os.path.join(path,'Data_template')
            f_list = [f for f in os.listdir(path1) if isfile(join(path1, f))]
            f_list1 = [f for f in os.listdir(path2) if isfile(join(path2, f))]
        return render_template("Lat.html",project=project,file_list = f_list,file_list1 = f_list1)
    if request.method == 'POST':
        try:
            req = request.get_json()
            df=pd.DataFrame(req)
            df.columns = ["Sno","Business_Rule","String_Name"]
            db_store(project,dataframe=df,uploads='template')
            controller(project)
            print(os.getcwd())
            path=os.getcwd()
            path=os.path.join(path,'Data')
            path=os.path.join(path,project)
            path1=os.path.join(path,'Data_defination')
            path2=os.path.join(path,'Data_template')
            f_list = [f for f in os.listdir(path1) if isfile(join(path1, f))]
            f_list1 = [f for f in os.listdir(path2) if isfile(join(path2, f))]
            return redirect(f'/{project}',file_list = f_list,file_list1 = f_list1)
        except Exception as e:
            f_list = [f for f in os.listdir(path1) if isfile(join(path1, f))]
            f_list1 = [f for f in os.listdir(path2) if isfile(join(path2, f))]
            return redirect(f'/{project}',file_list = f_list,file_list1 = f_list1)
@app.route('/function', methods=['POST'])
def get_ses():
    print("server has called api")
    text = request.form['text']
    #https://docs.google.com/spreadsheets/d/1Gztm9o8JEPibPWEDwH54qBG5kwj51ILDOk_dxK6uTSY/export?format=xlsx&gid=1330027783
    print(text)
    try:
        df=pd.read_excel(text)
    except Exception as e:
        Exceptionerror="Please check the url once"
        print(Exceptionerror)
        return render_template('upload.html',Exceptionerror=Exceptionerror)
    return render_template('upload0.html',tables=[df.to_html(index=False, classes=' table-hover table-condensed  table-striped center')])






@app.route('/<project>/sortby', methods = ['GET','POST'])
def sortby(project):

    if request.method == 'GET':

        dbname = project
        conn = sqlite3.connect(os.path.join(os.getcwd(),"Data",dbname+'.db'))
        f=[','.join(x) for x in conn.execute("SELECT name FROM sqlite_master WHERE type='table';").fetchall()]
        f=[x for x in f if x!='template']
        b = {}

        if f!=[]:
         for db in f:
          df=pd.read_sql(f'select * from {db}',conn).drop('index',axis=1)

          b[db] = df.columns.tolist()


         sheet = [(k, v) for k, v in b.items()]
         if os.path.exists(os.path.join(os.getcwd(),'Data',project+'_sort.json')):
             with open(os.path.join(os.getcwd(),'Data',project+'_sort.json')) as json_file:
                 proj_json = json.load(json_file)

         return render_template("Sortby.html",project=project,sheets=sheet,sheets_selected=list(proj_json.keys()),columns_selected=proj_json.values())
        else:
         return render_template("Sortby.html",project=project,sheets=[],sheets_selected=[],columns_selected=[])



    if request.method == 'POST':




        d=dict([ [(x,a) for a in request.form.getlist(x) if a!='on'][0] for x in [x[0] for x in [tup for tup in request.form.items() if any(i in tup for i in ['on'])]]])

        with open(os.path.join(os.getcwd(),'Data',project+'_sort.json'),'w') as json_file:
            json.dump(d, json_file)





        dbname = project
        conn = sqlite3.connect(os.path.join(os.getcwd(),"Data",dbname+'.db'))
        f=[','.join(x) for x in conn.execute("SELECT name FROM sqlite_master WHERE type='table';").fetchall()]
        f=[x for x in f if x!='template']
        b = {}
        if f!=[]:
         for db in f:
          df=pd.read_sql(f'select * from {db}',conn).drop('index',axis=1)
          #b["SheetName"] = db
          b[db] = df.columns.tolist()
          #b["ColumnList"] = df.columns.tolist()
          #a = {"SheetName":db,"ColumnList":df.columns.tolist()}
          #b = json.loads(a).update(a)/


         sheet = [(k, v) for k, v in b.items()]
        controller(project)

        with open(os.path.join(os.getcwd(),'Data',project+'_sort.json')) as json_file:
                 proj_json = json.load(json_file)

        return render_template("Sortby.html",project=project,sheets=sheet,sheets_selected=list(proj_json.keys()),columns_selected=list(proj_json.values()))



@app.route('/<project>/<uploads>', methods = ['POST'])
def upload_input(project,uploads):



   if request.method == 'POST':
    try:
      #f = request.files['file']

      uploaded_files = request.files

      uploaded_files = uploaded_files.to_dict(flat=False)

      files = uploaded_files["file"]


      if set([file.filename for file in files])!={''}:
       for file in files:
          if (file.filename).split(".")[1] in ALLOWED_DOCS:
              file.save(os.path.join('temp',secure_filename(file.filename)))


              db_store(project,filename=file.filename,uploads=uploads)

              controller(project)


          else:

              return redirect(f'/{project}')


       flash('File has been uploaded successfully')
       return redirect(f'/{project}')
      else:
       return redirect(f'/{project}')
    except Exception as e:

      return redirect(f'/{project}')


#Home page
@app.route('/', methods = ['GET','POST'])
def upload_file():
   if request.method=='GET' :
       df=pd.DataFrame([x.split('.')[0] for x in os.listdir('Data') if 'favicon' not in x and '.db' in x])

       if not df.empty:
        df.columns=['columns']
        df['sno']=['val'+str(x) for x in np.arange(len(df))+1]
        df['links']=request.host_url+df['columns']

        df=df.replace(np.nan,'',regex=True).drop_duplicates()
        items = df.to_dict('records')

       else:
         items={}
       if os.path.exists(os.path.join('Data','logo.png')):
          logo='yes'
       else:
         logo=''
       return render_template('Multi.html',items=items,logo=logo)

   if request.method=='POST':

    jsonresponse=request.form.to_dict()


    if jsonresponse=={}:

       flash('Nothing has been selected')
       return redirect('/')
    elif 'textbox' not  in jsonresponse.keys() :
     try:
      l=[x[1] for x in list(jsonresponse.items()) if 'chk[]' not in x[0]]

      for i in l:
       controller(i)
      document = Document()
      if os.path.exists(os.path.join(os.getcwd(),'Data','logo.png')):
       document.add_picture(os.path.join(os.getcwd(),'Data','logo.png'))
       last_paragraph = document.paragraphs[-1]
       last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

      #Project json
      with open(os.path.join(os.getcwd(),'Data','Project.json')) as json_file:
         json_decoded_consolidated = json.load(json_file)
      #run = document.add_paragraph().add_run()

      style = document.styles['Normal']
      font = style.font
      font.name = 'Calibri'
      font.size = Pt(28)
      font.color.rgb = RGBColor(0,0,0)
      font.bold = True
      paragraph1 = document.add_paragraph(json_decoded_consolidated['Projectheader']+"\n")
      #run = paragraph1.add_run()
      #font = run.font
      #font.name = 'Calibri'
      #font.size = Pt(20)
      #font.color.rgb = RGBColor(0,0,0)
      #font.bold = True
      paragraph_format = paragraph1.paragraph_format
      paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

      sentence = paragraph1.add_run(json_decoded_consolidated['Projectdescription']+"\n")
      sentence.font.size = Pt(20)
      sentence.font.color.rgb = RGBColor(192,192,192)



      for i in l:
       df = pd.read_csv(os.path.join(os.getcwd(),'Data',i+'_content.csv'))

       with open(os.path.join(os.getcwd(),'Data',i+'.json')) as json_file:
          json_decoded = json.load(json_file)
       #run = document.add_paragraph().add_run()

       style = document.styles['Normal']
       font = style.font
       font.name = 'Calibri'
       font.size = Pt(13)
       font.color.rgb = RGBColor(0,105,225)
       font.bold = True
       paragraph=document.add_paragraph(json_decoded['header']+"\n")
       sentence = paragraph.add_run(json_decoded['description']+"\n")
       sentence.font.size=Pt(11)
       sentence.font.italic = True
       for index,row in df.iterrows():

         sentence = paragraph.add_run(row.vals+"\n")
         sentence.font.name = 'Calibri'
         sentence.font.bold = False
         sentence.font.size = Pt(11)
         sentence.font.color.rgb=RGBColor(0,0,0)





      document.save(os.path.join(os.getcwd(),'temp','Combo.docx'))

      return send_from_directory('temp','Combo.docx',as_attachment=True,attachment_filename=f'Combo_{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}.docx')
     except Exception as e:
      flash('File could not be downloaded')
      return redirect('/')
    else:
     if jsonresponse['textbox'].lower() not in [x.split('.')[0].lower() for x in os.listdir('Data') if 'favicon' not in x]:

      return redirect(f"{request.host_url}{jsonresponse['textbox']}")
     else:
       flash('Project already exists with the same name')
       return redirect('/')

#download project level output here
@app.route('/<project>/output',methods = ['GET'])
def download_file(project):
    if request.method == 'GET':
     try:



       if not os.path.exists(os.path.join(os.getcwd(),'Data',project+'_content.csv')):

        controller(project)

       df=pd.read_csv(os.path.join(os.getcwd(),'Data',project+'_content.csv'))

       document=Document()


       with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
         json_decoded = json.load(json_file)
       #run = document.add_paragraph().add_run()

       style = document.styles['Normal']
       font = style.font
       font.name = 'Calibri'

       font.size = Pt(13)
       #font.color.rgb = RGBColor(0x42, 0x24, 0xE9)
       font.color.rgb = RGBColor(0,105,225)
       font.bold = True

       #download_format = json_decoded['format']
       paragraph = document.add_paragraph(json_decoded['header']+'\n')
       sentence = paragraph.add_run(json_decoded['description']+"\n")
       sentence.font.size=Pt(11)
       sentence.font.italic = True
       for index,row in df.iterrows():

        sentence = paragraph.add_run(row.vals+"\n")
        sentence.font.bold = False
        sentence.font.name = 'Calibri'
        sentence.font.size = Pt(11)
        sentence.font.color.rgb=RGBColor(0,0,0)


       document.save(os.path.join(os.getcwd(),'temp',f'{project}.docx'))
       #convert(os.path.join(os.getcwd(),'temp',f'{project}.docx'),os.path.join(os.getcwd(),'temp',f'{project}.pdf'))
       #if download_format == '#word':
       return send_from_directory('temp',f'{project}.docx',as_attachment=True,attachment_filename=f'{project}_{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}.docx')
       #elif download_format == '#pdf':
       #return send_from_directory('temp',f'{project}.pdf',as_attachment=True,attachment_filename=f'{project}_{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}.pdf')

     except Exception as e:

       return redirect(f'/{project}')

@app.route('/<project>/processdatatemplate', methods = ['GET','POST'])
def download_template(project):
    path=os.getcwd()
    path=os.path.join(path,'Data')
    path=os.path.join(path,project)
    path1=os.path.join(path,'Data_defination')
    path2=os.path.join(path,'Data_template')
    f_list = [f for f in os.listdir(path1) if isfile(join(path1, f))]
    f_list1 = [f for f in os.listdir(path2) if isfile(join(path2, f))]

    if (len(f_list)>=1 and len(f_list1)>=1):

        print("Project name:",project)
        va(project)
        return redirect(url_for('process', project=project))
    else:
        print('**********************----------------------',len(f_list))
        flash(u'Upload both files','error')
        return redirect(f'/{project}')
@app.route('/<project>/processed', methods = ['POST'])
def processed(project):
    path=os.getcwd()
    path=os.path.join(path,'Data')
    path=os.path.join(path,project)
    path1=os.path.join(path,project)
    pathf=os.path.join(path,'Updated_rules.csv')
    path1=os.path.join(path,'Data_defination')
    path2=os.path.join(path,'Data_template')
    f_list = [f for f in os.listdir(path1) if isfile(join(path1, f))]
    f_list1 = [f for f in os.listdir(path2) if isfile(join(path2, f))]
    if os.path.exists(pathf):
        path=os.getcwd()
        path=os.path.join(path,'Data')
        path=os.path.join(path,project)
        print('************************Calling************************')
        exportToWord(path)
        return redirect(url_for('process', project=project))
    elif (len(f_list)>=1 and len(f_list1)>=1):
        print("Project name:",project)
        return redirect(url_for('download_template', project=project))
    else:
        print('**********************----------------------',len(f_list))
        flash(u'Upload both files','error')
        return redirect(f'/{project}')
@app.route('/<project>/processdatatemplate1', methods = ['GET','POST'])
def download_template1(project):
    path=os.getcwd()
    path=os.path.join(path,'Data')
    path=os.path.join(path,project)
    pathf=os.path.join(path,'Updated_rules.csv')
    path1=os.path.join(path,'Data_defination')
    path2=os.path.join(path,'Data_template')
    f_list = [f for f in os.listdir(path1) if isfile(join(path1, f))]
    f_list1 = [f for f in os.listdir(path2) if isfile(join(path2, f))]
    if os.path.exists(pathf):
        df=pd.read_csv(pathf)
        return render_template('upload0.html', project=project,df1=df,tables=[df.to_html(index=False, classes=' table-hover table-condensed  table-striped center')])
    elif (len(f_list)>=1 and len(f_list1)>=1):
        print("Project name:",project)
        df=va1(project)
        return render_template('upload0.html', project=project,df1=df,tables=[df.to_html(index=False, classes=' table-hover table-condensed  table-striped center')])
        #return redirect(url_for('process', project=project))
    else:
        print('**********************----------------------',len(f_list))
        flash(u'Upload both files','error')
        return redirect(f'/{project}')
    #return redirect(f'/{project}')
@app.route('/<project>/change/<index>', methods = ['GET','POST'])
def changetemplate(project,index):
    print('Here**********************************************')
    path=os.getcwd()
    path1=os.path.join(path,'Data')
    path=os.path.join(path1,project)
    pathfile=os.path.join(path,'Updated_rules.csv')
    try:
        df=pd.read_csv(pathfile)
        text=request.form[index]
        print('reading file********************************************************')
        print(df.loc[int(index)-1,'updated_content'])
        df.loc[int(index)-1,'updated_content']=text
        print(df.loc[int(index)-1,'updated_content'])
        df.to_csv(pathfile,index=False)
        exportToWord(path1)
        return redirect(f'/{project}')
    except Exception as e:
        return redirect(f'/{project}')

@app.route('/<project>/change', methods = ['GET','POST'])
def changetemplate1(project):
    print('Heere**********************************************')
    path=os.getcwd()
    path1=os.path.join(path,'Data')
    path=os.path.join(path1,project)
    pathfile=os.path.join(path,'Updated_rules.csv')
    try:
        df=pd.read_csv(pathfile)
        total_rows = df.shape[0]
        print(total_rows)
        for i in range(total_rows):
            print(i)
            text=request.form[str(int(i+1))]
            print(text)
            if text=='nan':
                text=None
            df.loc[i,'updated_content']=text
        df.to_csv(pathfile,index=False)
        exportToWord(path1)
        return redirect(f'/{project}')
    except Exception as e:
        return redirect(f'/{project}')

    if (len(f_list)>=1 and len(f_list1)>=1):

        print("Project name:",project)
        df=va1(project)
        return render_template('upload0.html', df1=df,tables=[df.to_html(index=False, classes=' table-hover table-condensed  table-striped center')])
        #return redirect(url_for('process', project=project))
    else:
        print('**********************----------------------',len(f_list))
        flash(u'Upload both files','error')
        return redirect(f'/{project}')

@app.route('/<project>/description', methods = ['POST'])
def add_header(project):
    try:
     with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
      json_decoded = json.load(json_file)
    except:
     json_decoded={}

    json_decoded['header'] = request.form['Header']
    json_decoded['description']=request.form['Description']


    with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
     json.dump(json_decoded, json_file)

    return redirect(f'/{project}')

@app.route('/projectdescription', methods = ['POST'])
def add_project_header():
    try:
     with open(os.path.join(os.getcwd(),'Data','project.json')) as json_file:
      json_decoded = json.load(json_file)
    except:
     json_decoded={}

    json_decoded['Projectheader'] = request.form['ProjectHeader']
    json_decoded['Projectdescription']=request.form['ProjectDescription']


    with open(os.path.join(os.getcwd(),'Data','Project.json'),'w') as json_file:
     json.dump(json_decoded, json_file)

    return redirect(f'/')

@app.route('/logo', methods = ['POST'])
def upload_logo():

   if request.method == 'POST':
    try:
      #f = request.files['file']

      uploaded_files = request.files

      uploaded_files = uploaded_files.to_dict(flat=False)

      files = uploaded_files["file"]
      if set([file.filename for file in files])!={''}:
       for file in files:
          if (file.filename).split(".")[1].lower() in ALLOWED_IMAGES:
              file.save(os.path.join('Data',secure_filename('logo.png')))

      flash('The logo has been uploaded ')
      return redirect('/')
    except Exception as e:

     return redirect('/')


@app.route('/delete_logo',methods=['GET'])
def delete_logo():

 if os.path.exists(os.path.join('Data','logo.png')):

   os.remove(os.path.join('Data','logo.png'))
   flash('The logo has been deleted successfully')
   return redirect('/')
 else:
  flash('There is no logo present')
  return redirect('/')







@app.route('/<project>/preview', methods = ['GET','POST'])
def preview_file(project):
    if request.method=='GET':
      if not os.path.exists(os.path.join(os.getcwd(),'Data',project+'_content.csv')):

        controller(project)
      with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
          json_decoded = json.load(json_file)
      df=pd.read_csv(os.path.join(os.getcwd(),'Data',project+'_content.csv'))

      return {"Header":json_decoded['header'],"Description":json_decoded['description'],"Content":'\n'.join(df['vals'].tolist() )}
    if request.method=='POST':
       a = request.form.to_dict()
       with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
          json_decoded = json.load(json_file)

       json_decoded['header']=a['Header']
       json_decoded['description']=a['Description']
       with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
            json.dump(json_decoded, json_file)
       pd.DataFrame(a['Content'].split('\n'),columns=['vals']).to_csv(os.path.join(os.getcwd(),'Data',project+'_content.csv'))

       return redirect(f'/{project}')
@app.route('/<project>/rowbinder', methods = ['POST'])
def add_rowbinder(project):

    try:
     with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
      json_decoded = json.load(json_file)
    except Exception as e:

     json_decoded={}

    json_decoded['rowbinder'] = request.form['rowbinder']



    with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
     json.dump(json_decoded, json_file)

    controller(project)
    return redirect(f'/{project}')

######################################################################
############################ New Enhancements#########################
######################################################################

#uplpoad files here
@app.route('/Resolution/<uploads>', methods = ['POST'])
def upload_ResolutionFiles(uploads):
    if request.method == 'POST':

        try:

            uploaded_files = request.files.to_dict(flat=False)["file"]
            #uploaded_files = uploaded_files.to_dict(flat=False)
            #files = uploaded_files["file"]

            for file in uploaded_files:

                #if file.filename.split('.')[-1] == 'docx':

                #    print('hey',file.filename)
                #elif file.filename.split('.')[-1] == 'xlsx':
                #    print('hi',file.filename)

                if uploads == 'word':
                    print('hi')
                    file.save(os.path.join('temp',secure_filename('word.docx')))
                    #doc = docx.Document(os.path.join(os.getcwd(),'sample.xml'))
                    print('2')
                    doc = docx.Document(os.path.join(os.getcwd(),'temp','word.docx'))
                    print('1')
                    fullText = []
                    for para in doc.paragraphs:
                        fullText.append(para.text)
                    print(len(fullText))

            return redirect('/')

        except Exception as e:
            print(str(e))
            return redirect('/')



@app.route('/<project>/templatedownload',methods = ['GET'])
def templatedown_file(project):
    if request.method == 'GET':
     try:

       df=db_read(project,'template')

       df.to_csv(os.path.join(os.getcwd(),'temp',f'{project}_template.csv'),index=False)

       return send_from_directory('temp',f'{project}_template.csv',as_attachment=True,attachment_filename=f'{project}_template_{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}.csv')

     except Exception as e:

       return redirect(f'/{project}')
@app.route('/<project>/configure', methods = ['POST'])
def add_config(project):

    try:
     with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
      json_decoded = json.load(json_file)

    except Exception as e:

     json_decoded={}

    json_decoded['key'] = request.form['Paragraph_Config']



    with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
     json.dump(json_decoded, json_file)

    controller(project)
    return redirect(f'/{project}')




@app.route('/<project>/format', methods = ['POST'])
def add_format(project):

    try:
     with open(os.path.join(os.getcwd(),'Data',project+'.json')) as json_file:
      json_decoded = json.load(json_file)
    except Exception as e:

     json_decoded={}

    json_decoded['format'] = request.form['format']



    with open(os.path.join(os.getcwd(),'Data',project+'.json'),'w') as json_file:
     json.dump(json_decoded, json_file)

    #controller(project)
    return redirect(f'/{project}')

if __name__ == '__main__':
   app.run(host='0.0.0.0',port=5000,debug=True)
   #serve(app,host='0.0.0.0',port=5000)
