import pandas as pd,numpy as np,xlrd
from tqdm import tqdm
from math import *
import os
from os.path import isfile, join
from docx import Document
from docx.shared import Inches
import pandas as pd
ctr=1
optionaltext=[]

from docx import Document
from docx.shared import Inches
import pandas as pd
from functools import reduce
import operator
from docx.enum.text import WD_COLOR_INDEX


def exportToWord(path):
   print('***********************************************************************')
   print(path)
   df=pd.read_csv(os.path.join(path,'Updated_rules.csv'))
   df=df[df['updated_content'].notnull()]

   document = Document()
   def addWithHighlight(c,para):
        sentence=para.add_run(c)
        sentence.font.highlight_color=WD_COLOR_INDEX.YELLOW


   def customHighlight(c,para):

        a=reduce(operator.iconcat,[ [('{'+x[0]+'}','h'),(x[1],'n')] if type(x)==list else [(x,'n')] for x in [y.split('}') if '}' in y else y for y in  c.split('{')]],[])

        for i in a:
         if i[1]=='h':
          sentence=para.add_run(i[0])
          sentence.font.highlight_color=WD_COLOR_INDEX.GRAY_50
         else:
             sentence=para.add_run(i[0])





   def addWithoutHighlight(c,para):
        sentence=para.add_run(c)
        sentence.font.highlight_color=WD_COLOR_INDEX.YELLOW
   def add_heading(b,d,c):
    if b in ['Cat - Component - Heading']:
        document.add_heading(c, 0)
    elif ((b in ['Cat - Paragraph','Cat - Paragraph - GG']) and (d =='Optional')):
        para=document.add_paragraph()
        addWithHighlight(c,para)
    elif ((b in ['Cat - Paragraph','Cat - Paragraph - GG']) and (d =='Required')):
        para=document.add_paragraph()
        customHighlight(c,para)
        print('Required Paragraph added ')
    else:
        print('No criteria to generate doc')

   df.apply(lambda x: add_heading(x['asset type'],x['usage'],x['updated_content']),axis=1)
   document.save(os.path.join(path,'Document.docx'))





def va(project):
    path=os.getcwd()
    path=os.path.join(path,'Data')
    path=os.path.join(path,project)
    path1=os.path.join(path,'Data_defination')
    path2=os.path.join(path,'Data_template')
    f_list = [f for f in os.listdir(path1) if isfile(join(path1, f))]
    f_list1 = [f for f in os.listdir(path2) if isfile(join(path2, f))]
    f_list=f_list[0]
    f_list1=f_list1[0]

    df=pd.read_excel(os.path.join(path2,f_list1))
    df.columns=[x.lower() for x in df.columns]
    df=df.replace(np.nan,'',regex=True)

    a={}
    ##reading sheets into dictionary of dataframes

    xls = xlrd.open_workbook(os.path.join(path1,f_list), on_demand=True)
    a={sheet:pd.read_excel(f_list,sheet_name=sheet) for sheet in xls.sheet_names()}

    collist={x[0]:x[1].columns.tolist() for x in a.items()}
    con_list=[]
    [ con_list.extend(i) for i in collist.values()]
    ctr=1

    def evaluation(val):

    ##logic to identify the dataframe where the columns are present


        if val!='':

            val=val.replace('\n','')
            st=[]
            for i in con_list:
                vals=val.replace(f'<{i}>',f'{{{i}}}')
                if val!=vals:
                    st.append(i)
                val=vals


            sb={}

            for i in st:
                for j in collist.items():
                    if i in j[1]:
                        sb[i]=j[0]



            sc={i[0]:a[i[1]].loc[0,i[0]] for  i in sb.items()}


            val=val.format(**sc)
            val=val.replace(' & ',' and ').replace(' | ',' or ')

            return  eval(val)
            #exec(f'output={val}')
            #print(output)
        # return output
        else:
            return True

        asset_types=df['asset type'].drop_duplicates().tolist()



        def f(cmd,content,type):

            if cmd=='':
                return True

            else:
                return evaluation(cmd)







    tqdm.pandas()

    #df['eval_result']=df.progress_apply(lambda x:f(x['calculation rule'],x.content,x['asset type']),axis=1)



    def flag_calculate(a,b,c):
        if b in ['Cat - Component - Heading','Cat - IP','Cat - Component','Cat - Paragraph - GG']:
            return 'normal'
        elif b=='Cat - Paragraph':
            return 'probable'
        else:
            return ''





    df['flag']=df.progress_apply(lambda x: flag_calculate(x['calculation rule'],x['asset type'],x['content']),axis=1)


    def flag_calculate(a,b):
        if a=='probable':
            if '@FOREACHROW' in b:
                return 'loop'
            elif '@FOR<' in b:
                return 'list'
            else:
                return 'resolve'


        else:
            return a



    df['flag']=df.progress_apply(lambda x: flag_calculate(x['flag'],x['content']),axis=1)




    def eval_calculate(a,b):
        if b  =='normal':
            return evaluation(a)
        elif b=='resolve':
            return evaluation(a)
        else:
            return ''





    df['eval']=df.progress_apply(lambda x: eval_calculate(x['calculation rule'],x['flag']),axis=1)

    df=df.replace(np.nan,'',regex=True)



    def content_update(val,iterations):
        st=[]
        for i in con_list:
            vals=val.replace(f'<{i}>',f'{{{i}}}')
            if val!=vals:
                st.append(i)
                val=vals

        sb={}

        for i in st:
            for j in collist.items():
                if i in j[1]:
                    sb[i]=j[0]

        if type(iterations)==int:

            sc={i[0]:a[i[1]].loc[iterations,i[0]] for  i in sb.items()}


            for i in sc.items():

                val=val.replace(f'{{{i[0]}}}',str(i[1]))
            return val
        if type(iterations)==list:
            o=[]
            for index in iterations:
                sc={i[0]:a[i[1]].loc[index,i[0]] for  i in sb.items()}

                val1=val
                for i in sc.items():

                    val1=val1.replace(f'{{{i[0]}}}',str(i[1]))

                o.append(val1)


            o=','.join([a for a in o if a!=o[-1]])+' and '+o[-1]

            return o

    def loop_update(val):

        val=val.strip('@END')
        col=val.split('@FOREACHROW')[1].split('<<')[1].split('>>')[0]
        for j in collist.items():
            if col in j[1]:
                sb=j[0]

        return  content_update(val.split('@FOREACHROW')[1].split('>>')[1], list(range(len(a[sb][col]))))






    def f(a,b,c,d):
        if b=='resolve' and c==True:
            return content_update(a,iterations=0)
        elif d in ['Cat - Component - Heading','Cat - Paragraph - GG'] and c==True:
            return a

        elif b=='loop':
            return loop_update(a)

        else:
            return ''

    df['updated_content']=df.progress_apply(lambda x:f(x['content'],x.flag,x.eval,x['asset type']),axis=1)
    ctr=1
    def f(a):

        global ctr
        if a=='Cat - Component':
            ctr+=1
            return ctr
        else:
            return ctr




    df['component_index']=df.progress_apply(lambda x:f(x['asset type']),axis=1)



    df=df.loc[~df['component_index'].isin(df.loc[(df['asset type']=='Cat - Component') & (df['eval']==False)]['component_index'].tolist())]

    ###action


    if df.loc[df['asset type']=='Cat - IP','eval'].iloc[0]:
        filepath=os.path.join(path,'Updated_rules.csv')
        df.to_csv(filepath,index=False)
        print(path,filepath)
        exportToWord(path)


def va1(project):
    path=os.getcwd()
    path=os.path.join(path,'Data')
    path=os.path.join(path,project)
    path1=os.path.join(path,'Data_defination')
    path2=os.path.join(path,'Data_template')
    f_list = [f for f in os.listdir(path1) if isfile(join(path1, f))]
    f_list1 = [f for f in os.listdir(path2) if isfile(join(path2, f))]
    f_list=f_list[0]
    f_list1=f_list1[0]

    df=pd.read_excel(os.path.join(path2,f_list1))
    df.columns=[x.lower() for x in df.columns]
    df=df.replace(np.nan,'',regex=True)

    a={}
    ##reading sheets into dictionary of dataframes

    xls = xlrd.open_workbook(os.path.join(path1,f_list), on_demand=True)
    a={sheet:pd.read_excel(f_list,sheet_name=sheet) for sheet in xls.sheet_names()}

    collist={x[0]:x[1].columns.tolist() for x in a.items()}
    con_list=[]
    [ con_list.extend(i) for i in collist.values()]
    ctr=1

    def evaluation(val):

    ##logic to identify the dataframe where the columns are present


        if val!='':

            val=val.replace('\n','')
            st=[]
            for i in con_list:
                vals=val.replace(f'<{i}>',f'{{{i}}}')
                if val!=vals:
                    st.append(i)
                val=vals


            sb={}

            for i in st:
                for j in collist.items():
                    if i in j[1]:
                        sb[i]=j[0]



            sc={i[0]:a[i[1]].loc[0,i[0]] for  i in sb.items()}


            val=val.format(**sc)
            val=val.replace(' & ',' and ').replace(' | ',' or ')

            return  eval(val)
            #exec(f'output={val}')
            #print(output)
        # return output
        else:
            return True

        asset_types=df['asset type'].drop_duplicates().tolist()



        def f(cmd,content,type):

            if cmd=='':
                return True

            else:
                return evaluation(cmd)







    tqdm.pandas()

    #df['eval_result']=df.progress_apply(lambda x:f(x['calculation rule'],x.content,x['asset type']),axis=1)



    def flag_calculate(a,b,c):
        if b in ['Cat - Component - Heading','Cat - IP','Cat - Component','Cat - Paragraph - GG']:
            return 'normal'
        elif b=='Cat - Paragraph':
            return 'probable'
        else:
            return ''





    df['flag']=df.progress_apply(lambda x: flag_calculate(x['calculation rule'],x['asset type'],x['content']),axis=1)


    def flag_calculate(a,b):
        if a=='probable':
            if '@FOREACHROW' in b:
                return 'loop'
            elif '@FOR<' in b:
                return 'list'
            else:
                return 'resolve'


        else:
            return a



    df['flag']=df.progress_apply(lambda x: flag_calculate(x['flag'],x['content']),axis=1)




    def eval_calculate(a,b):
        if b  =='normal':
            return evaluation(a)
        elif b=='resolve':
            return evaluation(a)
        else:
            return ''





    df['eval']=df.progress_apply(lambda x: eval_calculate(x['calculation rule'],x['flag']),axis=1)

    df=df.replace(np.nan,'',regex=True)



    def content_update(val,iterations):
        st=[]
        for i in con_list:
            vals=val.replace(f'<{i}>',f'{{{i}}}')
            if val!=vals:
                st.append(i)
                val=vals

        sb={}

        for i in st:
            for j in collist.items():
                if i in j[1]:
                    sb[i]=j[0]

        if type(iterations)==int:

            sc={i[0]:a[i[1]].loc[iterations,i[0]] for  i in sb.items()}


            for i in sc.items():

                val=val.replace(f'{{{i[0]}}}',str(i[1]))
            return val
        if type(iterations)==list:
            o=[]
            for index in iterations:
                sc={i[0]:a[i[1]].loc[index,i[0]] for  i in sb.items()}

                val1=val
                for i in sc.items():

                    val1=val1.replace(f'{{{i[0]}}}',str(i[1]))

                o.append(val1)


            o=','.join([a for a in o if a!=o[-1]])+' and '+o[-1]

            return o

    def loop_update(val):

        val=val.strip('@END')
        col=val.split('@FOREACHROW')[1].split('<<')[1].split('>>')[0]
        for j in collist.items():
            if col in j[1]:
                sb=j[0]

        return  content_update(val.split('@FOREACHROW')[1].split('>>')[1], list(range(len(a[sb][col]))))






    def f(a,b,c,d):
        if b=='resolve' and c==True:
            return content_update(a,iterations=0)
        elif d in ['Cat - Component - Heading','Cat - Paragraph - GG'] and c==True:
            return a

        elif b=='loop':
            return loop_update(a)

        else:
            return ''

    df['updated_content']=df.progress_apply(lambda x:f(x['content'],x.flag,x.eval,x['asset type']),axis=1)
    ctr=1
    def f(a):

        global ctr
        if a=='Cat - Component':
            ctr+=1
            return ctr
        else:
            return ctr




    df['component_index']=df.progress_apply(lambda x:f(x['asset type']),axis=1)



    df=df.loc[~df['component_index'].isin(df.loc[(df['asset type']=='Cat - Component') & (df['eval']==False)]['component_index'].tolist())]

    ###action


    if df.loc[df['asset type']=='Cat - IP','eval'].iloc[0]:
        #os.remove(os.path.join(path,'Updated_rules.csv'))
        filepath=os.path.join(path,'Updated_rules.csv')
        df.to_csv(filepath,index=False)
        print(path,filepath)
        return df
