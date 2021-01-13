# -*- coding: utf-8 -*-
"""
Created on Tue Nov 24 21:24:46 2020

@author: nimondal
"""


from docx import Document
from docx.shared import Inches
import pandas as pd
from functools import reduce
import operator


def exportToWord(file):
   df=pd.read_csv(file)
   df=df[df['updated_content'].notnull()]
   from docx.enum.text import WD_COLOR_INDEX
   document = Document()
   def addWithHighlight(c,para):
        sentence=para.add_run(c)
        sentence.font.highlight_color=WD_COLOR_INDEX.YELLOW
        

   def customHighlight(c,para):
        a=reduce(operator.iconcat,[ [('{'+x[0]+'}','h'),(x[1],'n')] if type(x)==list else [(x,'n')] for x in [y.split('}') if '}' in y else y for y in  c.split('{')]],[])	
     
        for i in a:
         if i[1]=='h':
          sentence=para.add_run(i[0])
          sentence.font.highlight_color=WD_COLOR_INDEX.GRAY_50
         else:
             sentence=para.add_run(i[0])
             




   def addWithoutHighlight(c,para):
        sentence=para.add_run(c)
        sentence.font.highlight_color=WD_COLOR_INDEX.YELLOW     
   def add_heading(b,d,c):
    if b in ['Cat - Component - Heading']:
        document.add_heading(c, 0)
    elif ((b in ['Cat - Paragraph','Cat - Paragraph - GG','Cat - Paragraph - AG']) and (d =='Optional')):
        para=document.add_paragraph()
        addWithHighlight(c,para)
    elif ((b in ['Cat - Paragraph','Cat - Paragraph - GG','Cat - Paragraph - AG']) and (d =='Required')):
        para=document.add_paragraph(c)
        customHighlight(c,para)        
        print('Required Paragraph added ')
    else:
        print('No criteria to generate doc')
  
   df.apply(lambda x: add_heading(x['asset type'],x['usage'],x['updated_content']),axis=1)   
   document.save('demo.docx')

   document.save('demo.docx')
   print('Done')
   #print(optionaltext)

exportToWord('Updated_rules.csv')
